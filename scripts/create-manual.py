from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
ASSETS_DIR = DOCS_DIR / "assets"
OUTPUT = DOCS_DIR / "英単語1800_PWA_使い方マニュアル.docx"

APP_URL = "https://brestai20250522-bit.github.io/eitango1800/"

FONT = "Yu Gothic"
NAVY = RGBColor(12, 28, 46)
TEAL = RGBColor(0, 121, 111)
MUTED = RGBColor(82, 96, 112)
ROSE = RGBColor(190, 30, 75)
BLUE = RGBColor(43, 107, 215)
GOLD = RGBColor(171, 92, 0)
LIGHT_TEAL = "EAF7F5"
LIGHT_BLUE = "EEF4FF"
LIGHT_ROSE = "FFF1F5"
LIGHT_GOLD = "FFF8E8"
LINE = "D9E2EA"


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: float = 10.5, color: RGBColor = NAVY, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    # Named override: A4 is easier for Japanese classroom printing than US Letter.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Title", 24, NAVY, 0, 8),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 16, TEAL, 13, 7),
        ("Heading 2", 13, NAVY, 10, 5),
        ("Heading 3", 11.5, NAVY, 8, 3),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    for list_style in ["List Bullet", "List Number"]:
        style = styles[list_style]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.font.color.rgb = NAVY
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.16

    header = section.header.paragraphs[0]
    header.text = "英単語1800 PWA 使い方マニュアル"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(header, 8.5, MUTED, False)

    footer = section.footer.paragraphs[0]
    footer.text = "URL: " + APP_URL
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer, 8, MUTED, False)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ["top", "left", "bottom", "right"]:
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 120, start: int = 140, bottom: int = 120, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_fixed_width(table, widths_inches: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            row.cells[idx].width = Inches(width)


def add_kicker(doc: Document, text: str, color: RGBColor = TEAL) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, 9.5, color, True)


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    set_paragraph_font(p, 24, NAVY, True)
    if subtitle:
        sub = doc.add_paragraph(style="Subtitle")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.add_run(subtitle)
        set_paragraph_font(sub, 11.5, MUTED, False)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    set_paragraph_font(p, 16 if level == 1 else 13 if level == 2 else 11.5, TEAL if level == 1 else NAVY, True)


def add_body(doc: Document, text: str, bold: bool = False, color: RGBColor = NAVY) -> None:
    p = doc.add_paragraph()
    p.add_run(text)
    set_paragraph_font(p, 10.5, color, bold)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    set_paragraph_font(p, 10.2, NAVY, False)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    set_paragraph_font(p, 10.2, NAVY, False)


def add_callout(doc: Document, label: str, body: str, fill: str = LIGHT_TEAL, accent: RGBColor = TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_fixed_width(table, [6.9])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell)
    set_cell_margins(cell, 150, 180, 150, 180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(label)
    set_run_font(label_run, 10.5, accent, True)
    p.add_run("  ")
    body_run = p.add_run(body)
    set_run_font(body_run, 10, NAVY, False)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_mode_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table_fixed_width(table, [1.45, 2.7, 2.75])
    table.style = "Table Grid"
    headers = ["モード", "使う場面", "記録されるもの"]
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        shade_cell(cell, "E8EEF5")
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.add_run(text)
        set_paragraph_font(p, 9.5, NAVY, True)

    rows = [
        ("ランダム確認", "全範囲または単元を選んで小テスト。普段の確認に使う。", "正答率、苦手単語、復習予定に反映"),
        ("徹底特訓", "自分で選んだ苦手単語だけを何度も練習する。", "通常成績には反映しない"),
        ("今日の復習", "間違えた単語が時間をおいて出てきた時に解く。", "復習ステージと定着に反映"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_border(cells[i])
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.add_run(text)
            set_paragraph_font(p, 9.2, NAVY, i == 0)


def add_image(doc: Document, image_name: str, width: float, caption: str) -> None:
    image_path = ASSETS_DIR / image_name
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cap.add_run(caption)
    set_paragraph_font(cap, 8.5, MUTED, False)


def add_qr_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS_DIR / "app-qr.png"), width=Inches(1.75))
    url = doc.add_paragraph()
    url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url.add_run(APP_URL)
    set_paragraph_font(url, 8.5, MUTED, False)


def add_checklist_table(doc: Document, title: str, rows: list[tuple[str, str]], fill: str) -> None:
    add_heading(doc, title, 2)
    table = doc.add_table(rows=0, cols=2)
    table_fixed_width(table, [1.7, 5.2])
    for label, detail in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell)
        shade_cell(cells[0], fill)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = cells[0].paragraphs[0]
        p0.add_run(label)
        set_paragraph_font(p0, 9.5, NAVY, True)
        p1 = cells[1].paragraphs[0]
        p1.add_run(detail)
        set_paragraph_font(p1, 9.2, NAVY, False)


def build_manual() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)

    add_kicker(doc, "生徒・講師 共通版", TEAL)
    add_title(doc, "英単語1800 PWA 使い方マニュアル", "QRコードで開いて、4択テストと復習をすぐ始められます")
    add_qr_block(doc)
    add_callout(doc, "まずはこれだけ", "QRコードを読む → 「テストを実施する」 → モードを選ぶ。最初は「ランダム確認モード」がおすすめです。", LIGHT_TEAL, TEAL)
    add_heading(doc, "このアプリでできること", 1)
    add_mode_table(doc)
    add_callout(doc, "大切な注意", "成績や特訓メニューは、使っている端末の中に保存されます。別のスマホや別ブラウザには自動では引き継がれません。", LIGHT_GOLD, GOLD)

    doc.add_page_break()
    add_heading(doc, "生徒向け: 基本の使い方", 1)
    add_image(doc, "manual-home.png", 2.5, "ホーム画面: まず「テストを実施する」を押します")
    add_number(doc, "QRコードを読み取り、アプリを開きます。")
    add_number(doc, "ホーム画面の「テストを実施する」を押します。")
    add_number(doc, "「ランダム確認モード」または「徹底特訓モード」を選びます。")
    add_number(doc, "英単語を見て、日本語訳を4択から選びます。選ぶと自動で次の問題へ進みます。")
    add_callout(doc, "迷ったら", "学校や塾の小テスト練習は「ランダム確認モード」。自分だけの苦手単語練習は「徹底特訓モード」です。", LIGHT_BLUE, BLUE)

    doc.add_page_break()
    add_heading(doc, "ランダム確認モード", 1)
    add_image(doc, "manual-mode-select.png", 2.45, "モード選択: 通常の確認テストはランダム確認モード")
    add_body(doc, "ランダム確認モードは、通常の成績データに反映されるテストです。苦手単語、正答率、復習予定を残したいときに使います。")
    add_number(doc, "「テストを実施する」を押します。")
    add_number(doc, "「ランダム確認モード」を選びます。")
    add_number(doc, "範囲を選びます。「全単元」または学習したい単元を選べます。")
    add_number(doc, "問題数を選びます。10問、20問、または1単元全問を選べます。")
    add_number(doc, "「スタート」を押して解きます。")
    add_image(doc, "manual-random-setup.png", 2.45, "範囲と問題数を選んでスタート")
    add_callout(doc, "復習のしくみ", "間違えた単語だけが復習対象になります。一定時間後に「今日の復習」に出てきます。", LIGHT_TEAL, TEAL)

    doc.add_page_break()
    add_heading(doc, "徹底特訓モード", 1)
    add_body(doc, "徹底特訓モードは、どうしても覚えられない単語だけを自分で選んで練習するモードです。通常の成績データには反映されないので、安心して何度でも練習できます。")
    add_image(doc, "manual-training-menus.png", 2.45, "メニューA、B、Cを保存できます")
    add_number(doc, "「テストを実施する」から「徹底特訓モード」を選びます。")
    add_number(doc, "メニューA、B、Cのどれかで「編集」を押します。")
    add_number(doc, "表示する単元を切り替えながら、練習したい単語をタップして選びます。")
    add_number(doc, "「保存」を押します。")
    add_number(doc, "メニュー一覧に戻り、「実施」を押すと選んだ単語だけでテストできます。")
    add_image(doc, "manual-training-editor.png", 2.45, "単元をまたいで単語を選べます")
    add_callout(doc, "活用例", "メニューAは今週の苦手、メニューBはテスト前に最後まで残った単語、メニューCは過去に間違えた重要語のように使い分けできます。", LIGHT_ROSE, ROSE)

    doc.add_page_break()
    add_heading(doc, "ホーム画面に追加する方法", 1)
    add_checklist_table(
        doc,
        "iPhone / iPad",
        [
            ("1", "Safariでアプリを開きます。ChromeではなくSafariが分かりやすいです。"),
            ("2", "画面下の共有ボタンを押します。"),
            ("3", "「ホーム画面に追加」を選びます。"),
            ("4", "右上の「追加」を押すと、ホーム画面にアイコンができます。"),
        ],
        LIGHT_TEAL,
    )
    add_checklist_table(
        doc,
        "Android",
        [
            ("1", "Chromeでアプリを開きます。"),
            ("2", "右上の「︙」メニューを押します。"),
            ("3", "「ホーム画面に追加」または「アプリをインストール」を選びます。"),
            ("4", "確認画面で「追加」または「インストール」を押します。"),
        ],
        LIGHT_BLUE,
    )
    add_callout(doc, "開けない・画面が古いとき", "パソコンは Ctrl + F5 で強制更新。スマホはブラウザで再読み込み、またはホーム画面アイコンから開き直してください。", LIGHT_GOLD, GOLD)

    doc.add_page_break()
    add_heading(doc, "講師向け: 配布と運用", 1)
    add_bullet(doc, "生徒には1ページ目のQRコード、またはURLを共有します。")
    add_bullet(doc, "生徒は単語リストやアプリ本体を変更できません。公開ページを開いて使うだけです。")
    add_bullet(doc, "生徒の成績データは各端末に保存されます。講師側で全員分を自動集計する機能はありません。")
    add_bullet(doc, "徹底特訓モードは練習用です。正答率や苦手単語には反映されません。")
    add_bullet(doc, "単語の追加や修正は、管理者がExcelを更新して再公開する運用です。")
    add_checklist_table(
        doc,
        "授業での使い分け例",
        [
            ("小テスト", "ランダム確認モードで単元を指定し、10問または20問を実施します。"),
            ("復習日", "ホームの「今日の復習」を押し、間違えた単語だけを解き直します。"),
            ("個別対策", "徹底特訓モードで、生徒ごとに覚えにくい単語メニューを作らせます。"),
            ("定期前", "全単元ランダムで20問を繰り返し、抜けを確認します。"),
        ],
        LIGHT_TEAL,
    )
    add_heading(doc, "よくある質問", 2)
    add_bullet(doc, "通信料金は通常のWebページ閲覧程度です。OpenAI APIなどの利用料金はかかりません。")
    add_bullet(doc, "オフラインでも基本画面は開けますが、初回表示と更新時は通信が必要です。")
    add_bullet(doc, "プライベートブラウズや別端末では、進捗や特訓メニューが残らないことがあります。")
    add_bullet(doc, "同じ端末でもブラウザを変えると、保存データは別になります。")

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_manual()
